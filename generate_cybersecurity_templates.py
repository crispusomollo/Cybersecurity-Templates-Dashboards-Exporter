"""
Cybersecurity Templates & Dashboards Exporter (Executive Edition)
-----------------------------------------------------------------
Generates .docx, .xlsx, and .csv templates for all cybersecurity phases,
adds Excel dashboards, compresses output, generates a detailed executive summary,
and exports logs for reporting.

Author: GPT-5 (Generic Export Version)
"""

import os
import csv
import zipfile
import time
from datetime import datetime
import pandas as pd
from tqdm import tqdm
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.chart import PieChart, BarChart, LineChart, Reference
from openpyxl.utils import get_column_letter


# === CONFIGURATION ===
BASE_FOLDER = "Cybersecurity_Templates_&_Dashboards"
PHASES = [
    "Phase_1_Governance_Risk_SOC",
    "Phase_2_Technical_Controls_Monitoring",
    "Phase_3_Threat_Intel_Automation",
    "Phase_4_Compliance_Audit",
    "Phase_5_Awareness_Reporting"
]

SAMPLE_DATA = [
    ["Control ID", "Description", "Owner", "Status", "Last Reviewed"],
    ["CTRL-001", "Firewall configuration and rule audit", "Network Team", "Active", "2025-10-01"],
    ["CTRL-002", "User Access Review (Quarterly)", "IT Security", "In Progress", "2025-09-15"],
    ["CTRL-003", "Database Encryption Validation", "DBA", "Completed", "2025-10-20"],
    ["CTRL-004", "Vulnerability Scan Report", "SOC Team", "Completed", "2025-10-22"],
    ["CTRL-005", "Incident Response Plan Drill", "CISO Office", "Planned", "2025-11-01"]
]

TREND_DATA = [
    ["Month", "Compliance %"],
    ["June", 72],
    ["July", 76],
    ["August", 81],
    ["September", 85],
    ["October", 88],
    ["November", 92]
]


# === STEP 1: Create Folder Structure ===
def create_phase_folders():
    for phase in PHASES:
        os.makedirs(os.path.join(BASE_FOLDER, phase), exist_ok=True)
    os.makedirs(os.path.join(BASE_FOLDER, "Supporting_Files"), exist_ok=True)


# === STEP 2: Generate Per-Phase Files ===
def generate_phase_files():
    for phase in tqdm(PHASES, desc="Generating Phase Files", unit="phase"):
        phase_path = os.path.join(BASE_FOLDER, phase)

        # DOCX
        doc_path = os.path.join(phase_path, f"{phase}_Template.docx")
        doc = Document()
        doc.add_heading("Cybersecurity Phase Template", 0)
        doc.add_paragraph(
            "This document provides a structured cybersecurity implementation template for this phase. "
            "It includes objectives, responsibilities, and key controls."
        )
        doc.add_paragraph(f"Phase: {phase.replace('_', ' ')}")
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.save(doc_path)

        # XLSX
        xlsx_path = os.path.join(phase_path, f"{phase}_Dashboard.xlsx")
        wb = Workbook()

        # Sheet 1: Controls
        ws = wb.active
        ws.title = "Controls"
        for row in SAMPLE_DATA:
            ws.append(row)

        for col in ws.columns:
            max_len = max(len(str(cell.value)) if cell.value else 0 for cell in col)
            ws.column_dimensions[get_column_letter(col[0].column)].width = max_len + 2

        # Sheet 2: Status Summary
        statuses = [r[3] for r in SAMPLE_DATA[1:]]
        df = pd.Series(statuses).value_counts().reset_index()
        df.columns = ["Status", "Count"]
        ws2 = wb.create_sheet("Status Summary")
        ws2.append(["Status", "Count"])
        for _, row in df.iterrows():
            ws2.append(list(row))

        data_ref = Reference(ws2, min_col=2, min_row=1, max_row=len(df) + 1)
        labels_ref = Reference(ws2, min_col=1, min_row=2, max_row=len(df) + 1)

        pie = PieChart()
        bar = BarChart()
        pie.add_data(data_ref, titles_from_data=True)
        pie.set_categories(labels_ref)
        pie.title = "Control Status Distribution"
        bar.add_data(data_ref, titles_from_data=True)
        bar.set_categories(labels_ref)
        bar.title = "Control Status by Count"

        ws2.add_chart(pie, "D2")
        ws2.add_chart(bar, "D20")

        # Sheet 3: Compliance Trend
        ws3 = wb.create_sheet("Compliance Trend")
        for row in TREND_DATA:
            ws3.append(row)

        data_ref_trend = Reference(ws3, min_col=2, min_row=1, max_row=len(TREND_DATA))
        labels_ref_trend = Reference(ws3, min_col=1, min_row=2, max_row=len(TREND_DATA))

        line = LineChart()
        line.add_data(data_ref_trend, titles_from_data=True)
        line.set_categories(labels_ref_trend)
        line.title = "Monthly Compliance Trend"
        line.y_axis.title = "Compliance (%)"
        ws3.add_chart(line, "D4")

        wb.save(xlsx_path)

        # CSV
        csv_path = os.path.join(phase_path, f"{phase}_Data.csv")
        with open(csv_path, "w", newline="") as f:
            writer = csv.writer(f)
            writer.writerows(SAMPLE_DATA)

        time.sleep(0.4)


# === STEP 3: Create Summary Files ===
def generate_summary_report():
    summary = []
    for root, dirs, files in os.walk(BASE_FOLDER):
        if files and "Phase_" in root:
            phase_name = os.path.basename(root)
            total_size_kb = sum(os.path.getsize(os.path.join(root, f)) for f in files) / 1024
            summary.append({
                "Phase": phase_name,
                "Files": len(files),
                "Total_Size_KB": round(total_size_kb, 2),
                "Last_Updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })

    df_summary = pd.DataFrame(summary)
    summary_path = os.path.join(BASE_FOLDER, "Supporting_Files", "export_summary.csv")
    df_summary.to_csv(summary_path, index=False)

    # Executive Summary Word Report
    doc = Document()
    doc.add_heading("Cybersecurity Program – Executive Summary", 0)
    doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("This document provides an overview of all cybersecurity implementation phases, key metrics, and control status distribution.")

    doc.add_heading("Phase Overview", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Phase"
    hdr_cells[1].text = "Files"
    hdr_cells[2].text = "Total Size (KB)"
    hdr_cells[3].text = "Last Updated"

    for _, row in df_summary.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Phase"])
        row_cells[1].text = str(row["Files"])
        row_cells[2].text = str(row["Total_Size_KB"])
        row_cells[3].text = str(row["Last_Updated"])

    doc.add_paragraph("\nOverall compliance shows steady improvement across all phases, with most controls marked as 'Completed' or 'Active'.")
    doc.add_heading("Recommendations", level=1)
    doc.add_paragraph("- Continue quarterly user access reviews.")
    doc.add_paragraph("- Increase automation coverage for vulnerability management.")
    doc.add_paragraph("- Schedule an annual review of SOC incident response procedures.")
    doc.add_paragraph("- Maintain awareness and training programs for new staff.")

    doc_path = os.path.join(BASE_FOLDER, "Cybersecurity_Executive_Summary.docx")
    doc.save(doc_path)
    print(f"\n📘 Executive Summary generated at: {doc_path}")


# === STEP 4: ZIP Everything ===
def create_zip_archive():
    zip_filename = f"{BASE_FOLDER}.zip"
    with zipfile.ZipFile(zip_filename, "w", zipfile.ZIP_DEFLATED) as zipf:
        for root, _, files in os.walk(BASE_FOLDER):
            for file in files:
                zipf.write(os.path.join(root, file), os.path.relpath(os.path.join(root, file), BASE_FOLDER))
    print(f"✅ Export complete: {zip_filename}")


# === MAIN RUN ===
if __name__ == "__main__":
    print("🚀 Starting Cybersecurity Templates & Dashboard Export...\n")
    create_phase_folders()
    generate_phase_files()
    generate_summary_report()
    create_zip_archive()
    print("\n🎯 All files and executive report generated successfully.\n")

